from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import Any

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from engine.models.document_model import DocumentModel
from engine.models.paragraph_model import ParagraphModel
from engine.models.section_model import SectionModel
from engine.models.table_model import TableModel
from engine.shared.utils import twips_to_mm


@dataclass
class DocumentParser:
    """
    Парсер DOCX → первичная внутренняя модель.

    На этом шаге делаем:
    - последовательный обход тела документа (p/tbl),
    - извлечение текстов/стилей/форматирования,
    - базовую работу с секциями (поля, ориентация).
    """

    def parse_docx_bytes(self, docx_bytes: bytes) -> DocumentModel:
        doc = DocxDocument(BytesIO(docx_bytes))

        model = DocumentModel()
        model.numbering = self._extract_numbering_metadata(doc)

        # Стили (для отладки/эвристик).
        try:
            model.styles = {s.name: {} for s in doc.styles}
        except Exception:
            model.styles = {}

        # Секции и настройки полей.
        for idx, sec in enumerate(doc.sections):
            orientation = None
            try:
                orientation = str(sec.orientation)
            except Exception:
                orientation = None

            page_settings = {
                "orientation": orientation,
                "page_width_mm": twips_to_mm(sec.page_width.twips) if sec.page_width else None,
                "page_height_mm": twips_to_mm(sec.page_height.twips) if sec.page_height else None,
                "margins_mm": {
                    "left": twips_to_mm(sec.left_margin.twips),
                    "right": twips_to_mm(sec.right_margin.twips),
                    "top": twips_to_mm(sec.top_margin.twips),
                    "bottom": twips_to_mm(sec.bottom_margin.twips),
                },
            }
            model.page_settings[f"section_{idx}"] = page_settings
            model.sections.append(
                SectionModel(
                    index=idx,
                    orientation=orientation,
                    page_width_mm=page_settings["page_width_mm"],
                    page_height_mm=page_settings["page_height_mm"],
                    margin_left_mm=page_settings["margins_mm"]["left"],
                    margin_right_mm=page_settings["margins_mm"]["right"],
                    margin_top_mm=page_settings["margins_mm"]["top"],
                    margin_bottom_mm=page_settings["margins_mm"]["bottom"],
                )
            )

        # Проходим по телу документа в порядке следования.
        p_id = 0
        t_id = 0

        for child in doc.element.body:
            if not hasattr(child, "tag"):
                continue
            tag = child.tag

            if tag.endswith("}p") or tag.endswith("}pPr") or tag.endswith("p"):
                # child обычно именно w:p
                try:
                    dp = DocxParagraph(child, doc)
                except Exception:
                    continue

                text = (dp.text or "").strip()
                p_id += 1

                xml_element = getattr(dp, "_p", None)
                xml_str = xml_element.xml if xml_element is not None else ""

                starts_with_page_break = False
                if xml_str:
                    before_text = xml_str.split("<w:t", 1)[0]
                    starts_with_page_break = 'w:type="page"' in before_text

                # Runs metadata.
                runs_meta: list[dict[str, Any]] = []
                for run in dp.runs:
                    fs = run.font.size
                    font_size_pt = fs.pt if fs is not None else None
                    runs_meta.append(
                        {
                            "text": run.text,
                            "font_name": run.font.name,
                            "font_size_pt": font_size_pt,
                            "bold": bool(run.bold) if run.bold is not None else None,
                            "italic": bool(run.italic) if run.italic is not None else None,
                            "underline": bool(run.underline) if run.underline is not None else None,
                        }
                    )

                pf = dp.paragraph_format
                first_line_indent_cm = pf.first_line_indent.cm if pf.first_line_indent is not None else None
                left_indent_cm = pf.left_indent.cm if pf.left_indent is not None else None
                right_indent_cm = pf.right_indent.cm if pf.right_indent is not None else None
                space_before_pt = pf.space_before.pt if pf.space_before is not None else None
                space_after_pt = pf.space_after.pt if pf.space_after is not None else None
                line_spacing = pf.line_spacing

                alignment = pf.alignment
                if isinstance(alignment, WD_ALIGN_PARAGRAPH):
                    alignment_val = alignment
                else:
                    alignment_val = alignment

                style_name = dp.style.name if dp.style is not None else None
                list_info = self._extract_paragraph_list_info(dp, style_name)

                model.paragraphs.append(
                    ParagraphModel(
                        id=f"p{p_id}",
                        text=text,
                        style_name=style_name,
                        page_number=None,
                        section_index=0,
                        zone_type=None,
                        xml_ref=f"xml:p{p_id}",
                        alignment=str(alignment_val) if alignment_val is not None else None,
                        line_spacing=float(line_spacing) if line_spacing is not None else None,
                        first_line_indent_cm=first_line_indent_cm,
                        left_indent_cm=left_indent_cm,
                        right_indent_cm=right_indent_cm,
                        space_before_pt=space_before_pt,
                        space_after_pt=space_after_pt,
                        is_list_item=list_info["is_list_item"],
                        list_kind=list_info["list_kind"],
                        list_level=list_info["list_level"],
                        list_num_id=list_info["list_num_id"],
                        in_table=False,
                        in_header_footer=False,
                        xml_element=xml_element,
                        starts_with_page_break=starts_with_page_break,
                        runs_meta=runs_meta,
                    )
                )
                continue

            if tag.endswith("}tbl") or tag.endswith("tbl"):
                t_id += 1
                try:
                    dtbl = DocxTable(child, doc)
                except Exception:
                    continue

                table_model = TableModel(
                    id=f"t{t_id}",
                    zone_type=None,
                    page_number=None,
                    xml_ref=f"xml:tbl{t_id}",
                    rows=len(dtbl.rows),
                    cols=len(dtbl.columns),
                )
                model.tables.append(table_model)

                # Извлекаем ячейки как набор “вложенных параграфов” (для чекеров по шрифту).
                for r_idx, row in enumerate(dtbl.rows):
                    for c_idx, cell in enumerate(row.cells):
                        for pc_idx, para in enumerate(cell.paragraphs):
                            p_id += 1
                            text = (para.text or "").strip()
                            xml_element = getattr(para, "_p", None)
                            xml_str = xml_element.xml if xml_element is not None else ""
                            starts_with_page_break = False
                            if xml_str:
                                before_text = xml_str.split("<w:t", 1)[0]
                                starts_with_page_break = 'w:type="page"' in before_text

                            runs_meta: list[dict[str, Any]] = []
                            for run in para.runs:
                                fs = run.font.size
                                font_size_pt = fs.pt if fs is not None else None
                                runs_meta.append(
                                    {
                                        "text": run.text,
                                        "font_name": run.font.name,
                                        "font_size_pt": font_size_pt,
                                        "bold": bool(run.bold) if run.bold is not None else None,
                                        "italic": bool(run.italic) if run.italic is not None else None,
                                        "underline": bool(run.underline) if run.underline is not None else None,
                                    }
                                )

                            pf = para.paragraph_format
                            first_line_indent_cm = pf.first_line_indent.cm if pf.first_line_indent is not None else None
                            left_indent_cm = pf.left_indent.cm if pf.left_indent is not None else None
                            right_indent_cm = pf.right_indent.cm if pf.right_indent is not None else None
                            space_before_pt = pf.space_before.pt if pf.space_before is not None else None
                            space_after_pt = pf.space_after.pt if pf.space_after is not None else None
                            line_spacing = pf.line_spacing
                            alignment = pf.alignment
                            style_name = para.style.name if para.style is not None else None
                            list_info = self._extract_paragraph_list_info(para, style_name)

                            model.paragraphs.append(
                                ParagraphModel(
                                    id=f"t{t_id}_r{r_idx}_c{c_idx}_p{pc_idx}",
                                    text=text,
                                    style_name=style_name,
                                    page_number=None,
                                    section_index=0,
                                    zone_type=None,
                                    xml_ref=f"xml:p_{t_id}_{r_idx}_{c_idx}_{pc_idx}",
                                    alignment=str(alignment) if alignment is not None else None,
                                    line_spacing=float(line_spacing) if line_spacing is not None else None,
                                    first_line_indent_cm=first_line_indent_cm,
                                    left_indent_cm=left_indent_cm,
                                    right_indent_cm=right_indent_cm,
                                    space_before_pt=space_before_pt,
                                    space_after_pt=space_after_pt,
                                    is_list_item=list_info["is_list_item"],
                                    list_kind=list_info["list_kind"],
                                    list_level=list_info["list_level"],
                                    list_num_id=list_info["list_num_id"],
                                    in_table=True,
                                    in_header_footer=False,
                                    xml_element=xml_element,
                                    starts_with_page_break=starts_with_page_break,
                                    runs_meta=runs_meta,
                                )
                            )

        return model

    @staticmethod
    def _extract_paragraph_list_info(dp: DocxParagraph, style_name: str | None) -> dict[str, Any]:
        style_lower = (style_name or "").lower()
        style_is_list = "list" in style_lower or "спис" in style_lower
        list_kind = "numbered" if "number" in style_lower or "нумер" in style_lower else None
        list_level: int | None = None
        list_num_id: str | None = None

        p_pr = getattr(getattr(dp, "_p", None), "pPr", None)
        num_pr = getattr(p_pr, "numPr", None) if p_pr is not None else None
        if num_pr is not None:
            style_is_list = True
            list_kind = list_kind or "numbered"
            ilvl = getattr(num_pr, "ilvl", None)
            num_id = getattr(num_pr, "numId", None)
            if ilvl is not None:
                try:
                    list_level = int(ilvl.val)
                except Exception:
                    list_level = None
            if num_id is not None:
                try:
                    list_num_id = str(num_id.val)
                except Exception:
                    list_num_id = None

        return {
            "is_list_item": style_is_list,
            "list_kind": list_kind,
            "list_level": list_level,
            "list_num_id": list_num_id,
        }

    @staticmethod
    def _extract_numbering_metadata(doc) -> dict[str, Any]:
        sections: list[dict[str, Any]] = []
        for idx, sec in enumerate(doc.sections):
            footer_xml = sec.footer._element.xml if sec.footer is not None else ""
            first_footer_xml = sec.first_page_footer._element.xml if sec.first_page_footer is not None else ""
            sections.append(
                {
                    "index": idx,
                    "different_first_page": bool(sec.different_first_page_header_footer),
                    "footer_has_page_field": "PAGE" in footer_xml,
                    "first_footer_has_page_field": "PAGE" in first_footer_xml,
                }
            )

        numbering_xml = ""
        try:
            numbering_xml = doc.part.numbering_part.element.xml
        except Exception:
            numbering_xml = ""

        return {
            "sections": sections,
            "list_number_format_ok": DocumentParser._numbering_definitions_use_main_font(numbering_xml),
        }

    @staticmethod
    def _numbering_definitions_use_main_font(numbering_xml: str) -> bool:
        if not numbering_xml or "w:numFmt" not in numbering_xml:
            return False
        if 'w:val="decimal"' not in numbering_xml:
            return False
        return "Times New Roman" in numbering_xml and 'w:sz w:val="28"' in numbering_xml

