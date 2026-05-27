from __future__ import annotations

from dataclasses import dataclass, field
from typing import Protocol, Any

from io import BytesIO

from docx import Document as DocxDocument
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from engine.shared.constants import GOST_MONO_FONT_NAME

from engine.models.document_model import DocumentModel
from engine.models.violation_model import FixOperation, Violation


class BaseFixer(Protocol):
    """Интерфейс fixer-а: по нарушениям строит операции исправлений."""

    def supported_violation_codes(self) -> set[str]:  # pragma: no cover
        ...

    def build_fixes(self, document: DocumentModel, violations: list[Violation]) -> list[FixOperation]:  # pragma: no cover
        ...


@dataclass
class FixerConfig:
    # Приоритеты будут добавлены позже (структура, шрифт, подписи и т.д.)
    enable_conflict_resolution: bool = True


@dataclass
class Fixer:
    fixers: list[BaseFixer]
    config: FixerConfig = field(default_factory=FixerConfig)

    def build_operations(self, document: DocumentModel, violations: list[Violation]) -> list[FixOperation]:
        ops: list[FixOperation] = []
        for fixer in self.fixers:
            relevant = [v for v in violations if v.code in fixer.supported_violation_codes()]
            if not relevant:
                continue
            ops.extend(fixer.build_fixes(document, relevant))
        return ops

    def apply_fixes(self, docx_bytes: bytes, document: DocumentModel, operations: list[FixOperation]) -> tuple[bytes, list[FixOperation]]:
        if not operations:
            return docx_bytes, []

        # Приоритет: сначала страницы/поля, затем нумерация, затем шрифты.
        priority = {
            "SET_PAGE_MARGINS": 1,
            "SET_PAGE_NUMBERING": 2,
            "INSERT_EMPTY_PARAGRAPH_BEFORE": 15,
            "SET_FONT_FORMATTING": 20,
            "SET_FONT_FORMAT": 20,
            "SET_LIST_NUMBER_FORMAT": 21,
            "SET_PARAGRAPH_TEXT": 25,
            "SET_PARAGRAPH_FORMAT": 30,
        }
        operations_sorted = sorted(operations, key=lambda op: priority.get(op.action, 100))

        doc = DocxDocument(BytesIO(docx_bytes))

        insert_ops = [op for op in operations_sorted if op.action == "INSERT_EMPTY_PARAGRAPH_BEFORE"]
        other_ops = [op for op in operations_sorted if op.action != "INSERT_EMPTY_PARAGRAPH_BEFORE"]
        insert_positions = sorted(
            int(op.target_element_id[1:]) for op in insert_ops if op.target_element_id and op.target_element_id.startswith("p")
        )

        def _adjusted_element_id(element_id: str) -> str:
            if not element_id or not element_id.startswith("p"):
                return element_id
            n = int(element_id[1:])
            shift = sum(1 for pos in insert_positions if pos <= n)
            return f"p{n + shift}"

        def _find_paragraph_by_number(paragraph_number: int):
            from docx.text.paragraph import Paragraph as DocxParagraph

            p_id = 0
            for child in doc.element.body:
                if not hasattr(child, "tag"):
                    continue
                tag = child.tag
                if tag.endswith("}p") or tag.endswith("p"):
                    p_id += 1
                    if p_id == paragraph_number:
                        return DocxParagraph(child, doc)
            return None

        cumulative_insert = 0
        for op in sorted(insert_ops, key=lambda x: int(x.target_element_id[1:])):
            if not op.target_element_id:
                continue
            orig = int(op.target_element_id[1:])
            target_num = orig + cumulative_insert
            dp = _find_paragraph_by_number(target_num)
            if dp is not None:
                inserted = dp.insert_paragraph_before("")
                if op.meta:
                    self._apply_paragraph_format(inserted, op.meta)
                cumulative_insert += 1

        for op in other_ops:
            if op.action == "SET_PAGE_MARGINS":
                margins = (op.meta or {}).get("margins_mm") or {}
                for sec in doc.sections:
                    if "top" in margins:
                        sec.top_margin = Cm(margins["top"] / 10.0)
                    if "bottom" in margins:
                        sec.bottom_margin = Cm(margins["bottom"] / 10.0)
                    if "left" in margins:
                        sec.left_margin = Cm(margins["left"] / 10.0)
                    if "right" in margins:
                        sec.right_margin = Cm(margins["right"] / 10.0)

            elif op.action == "SET_PAGE_NUMBERING":
                number_from_second = bool((op.meta or {}).get("number_from_second_page", True))

                for sec in doc.sections:
                    sec.footer.is_linked_to_previous = False
                    sec.first_page_footer.is_linked_to_previous = False
                    sec.different_first_page_header_footer = True
                    if number_from_second:
                        self._clear_header_footer(sec.first_page_footer)
                    self._set_page_field_footer(sec.footer)

                self._enable_field_updates(doc)

            elif op.action == "INSERT_EMPTY_PARAGRAPH_BEFORE":
                continue

            elif op.action == "SET_PARAGRAPH_TEXT":
                target_id = _adjusted_element_id(op.target_element_id or "")
                new_text = op.value or ""
                alignment = (op.meta or {}).get("alignment")
                if not target_id:
                    continue

                p_id = 0
                t_id = 0
                found = False
                for child in doc.element.body:
                    if not hasattr(child, "tag"):
                        continue
                    tag = child.tag

                    if tag.endswith("}p") or tag.endswith("p"):
                        from docx.text.paragraph import Paragraph as DocxParagraph

                        dp = DocxParagraph(child, doc)
                        p_id += 1
                        element_id = f"p{p_id}"
                        if element_id == target_id:
                            self._set_paragraph_text(dp, new_text, alignment)
                            found = True
                            break

                    elif tag.endswith("}tbl") or tag.endswith("tbl"):
                        from docx.table import Table as DocxTable

                        t_id += 1
                        dtbl = DocxTable(child, doc)
                        for r_idx, row in enumerate(dtbl.rows):
                            for c_idx, cell in enumerate(row.cells):
                                for pc_idx, para in enumerate(cell.paragraphs):
                                    p_id += 1
                                    element_id = f"t{t_id}_r{r_idx}_c{c_idx}_p{pc_idx}"
                                    if element_id == target_id:
                                        self._set_paragraph_text(para, new_text, alignment)
                                        found = True
                                        break
                                if found:
                                    break
                            if found:
                                break
                    if found:
                        break

            elif op.action == "SET_PARAGRAPH_FORMAT":
                target_id = _adjusted_element_id(op.target_element_id or "")
                if not target_id:
                    continue
                meta = op.meta or {}
                p_id = 0
                t_id = 0
                found = False
                for child in doc.element.body:
                    if not hasattr(child, "tag"):
                        continue
                    tag = child.tag
                    if tag.endswith("}p") or tag.endswith("p"):
                        from docx.text.paragraph import Paragraph as DocxParagraph

                        dp = DocxParagraph(child, doc)
                        p_id += 1
                        element_id = f"p{p_id}"
                        if element_id == target_id:
                            self._apply_paragraph_format(dp, meta)
                            found = True
                            break
                    elif tag.endswith("}tbl") or tag.endswith("tbl"):
                        from docx.table import Table as DocxTable

                        t_id += 1
                        dtbl = DocxTable(child, doc)
                        for r_idx, row in enumerate(dtbl.rows):
                            for c_idx, cell in enumerate(row.cells):
                                for pc_idx, para in enumerate(cell.paragraphs):
                                    p_id += 1
                                    element_id = f"t{t_id}_r{r_idx}_c{c_idx}_p{pc_idx}"
                                    if element_id == target_id:
                                        self._apply_paragraph_format(para, meta)
                                        found = True
                                        break
                                if found:
                                    break
                            if found:
                                break
                    if found:
                        break

            elif op.action in {"SET_FONT_FORMATTING", "SET_FONT_FORMAT"}:
                main = (op.meta or {}).get("main") or {}
                listing = (op.meta or {}).get("listing") or {}
                listing_ids = set((op.meta or {}).get("listing_element_ids") or [])
                skip_ids = set((op.meta or {}).get("skip_element_ids") or [])
                listing_ids = {_adjusted_element_id(element_id) for element_id in listing_ids}
                skip_ids = {_adjusted_element_id(element_id) for element_id in skip_ids}

                # Нужно повторить нумерацию element_id, совпадающую с DocumentParser.
                # Поэтому делаем обход p/tbl и считаем p_id/ t_id так же.
                p_id = 0
                t_id = 0

                # Установка глобальных стилей “на нормальный” уменьшит риск наследования.
                try:
                    normal_style = doc.styles["Normal"]
                    normal_style.font.name = main.get("font_name")
                    normal_style.font.size = Pt(float(main.get("font_size_pt", 14)))
                except Exception:
                    pass

                for child in doc.element.body:
                    if not hasattr(child, "tag"):
                        continue
                    tag = child.tag

                    if tag.endswith("}p") or tag.endswith("p"):
                        from docx.text.paragraph import Paragraph as DocxParagraph

                        dp = DocxParagraph(child, doc)
                        p_id += 1
                        element_id = f"p{p_id}"

                        if not (dp.text or "").strip():
                            continue
                        if element_id in skip_ids:
                            continue
                        if element_id in listing_ids:
                            self._format_paragraph(dp, listing, is_listing=True)
                        else:
                            self._format_paragraph(dp, main, is_listing=False)

                    elif tag.endswith("}tbl") or tag.endswith("tbl"):
                        from docx.table import Table as DocxTable

                        t_id += 1
                        dtbl = DocxTable(child, doc)
                        for r_idx, row in enumerate(dtbl.rows):
                            for c_idx, cell in enumerate(row.cells):
                                for pc_idx, para in enumerate(cell.paragraphs):
                                    p_id += 1
                                    element_id = f"t{t_id}_r{r_idx}_c{c_idx}_p{pc_idx}"
                                    if not (para.text or "").strip():
                                        continue
                                    if element_id in skip_ids:
                                        continue
                                    if element_id in listing_ids:
                                        self._format_paragraph(para, listing, is_listing=True)
                                    else:
                                        self._format_paragraph(para, main, is_listing=False)

            elif op.action == "SET_LIST_NUMBER_FORMAT":
                self._format_list_numbering(doc, op.meta or {})

            else:
                # Неподдерживаемая операция.
                continue

        out = BytesIO()
        doc.save(out)
        return out.getvalue(), operations_sorted

    @staticmethod
    def _insert_empty_paragraph_before(dp: Any) -> None:
        from docx.oxml import OxmlElement

        new_p = OxmlElement("w:p")
        dp._p.addprevious(new_p)

    @staticmethod
    def _clear_paragraph_content(dp: Any) -> None:
        p_pr = dp._p.pPr
        for child in list(dp._p):
            if p_pr is not None and child is p_pr:
                continue
            dp._p.remove(child)

    @classmethod
    def _clear_header_footer(cls, part: Any) -> None:
        paragraphs = part.paragraphs or [part.add_paragraph()]
        for idx, paragraph in enumerate(paragraphs):
            cls._clear_paragraph_content(paragraph)
            if idx == 0:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    @classmethod
    def _set_page_field_footer(cls, footer: Any) -> None:
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        cls._clear_paragraph_content(paragraph)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run = paragraph.add_run()
        for element in cls._page_field_elements():
            run._r.append(element)

    @staticmethod
    def _page_field_elements() -> list[Any]:
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        fld_begin.set(qn("w:dirty"), "true")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "

        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")

        text = OxmlElement("w:t")
        text.text = "1"

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        return [fld_begin, instr_text, fld_separate, text, fld_end]

    @staticmethod
    def _enable_field_updates(doc: Any) -> None:
        settings = doc.settings._element
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")

    @staticmethod
    def _set_paragraph_text(dp: Any, new_text: str, alignment: str | None) -> None:
        for run in list(dp.runs):
            run.text = ""
        if dp.runs:
            dp.runs[0].text = new_text
        else:
            dp.add_run(new_text)

        if alignment:
            a = str(alignment).upper()
            if "CENTER" in a:
                dp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif "LEFT" in a:
                dp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif "RIGHT" in a:
                dp.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif "JUSTIFY" in a:
                dp.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    @staticmethod
    def _apply_paragraph_format(dp: Any, cfg: dict[str, Any]) -> None:
        alignment = cfg.get("alignment")
        indent_cm = cfg.get("first_line_indent_cm")
        line_spacing = cfg.get("line_spacing")
        left_indent_cm = cfg.get("left_indent_cm")
        right_indent_cm = cfg.get("right_indent_cm")
        space_before_pt = cfg.get("space_before_pt")
        space_after_pt = cfg.get("space_after_pt")
        font_name = cfg.get("font_name")
        font_size_pt = cfg.get("font_size_pt")
        bold = cfg.get("bold")
        italic = cfg.get("italic")
        underline = cfg.get("underline")

        if alignment:
            a = str(alignment).upper()
            if "CENTER" in a:
                dp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif "LEFT" in a:
                dp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif "RIGHT" in a:
                dp.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif "JUSTIFY" in a:
                dp.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        if indent_cm is not None:
            dp.paragraph_format.first_line_indent = Cm(float(indent_cm))
        if line_spacing is not None:
            dp.paragraph_format.line_spacing = float(line_spacing)
        if left_indent_cm is not None:
            dp.paragraph_format.left_indent = Cm(float(left_indent_cm))
        if right_indent_cm is not None:
            dp.paragraph_format.right_indent = Cm(float(right_indent_cm))
        if space_before_pt is not None:
            dp.paragraph_format.space_before = Pt(float(space_before_pt))
        if space_after_pt is not None:
            dp.paragraph_format.space_after = Pt(float(space_after_pt))

        for run in dp.runs:
            if font_name:
                run.font.name = font_name
                try:
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
                except Exception:
                    pass
            if font_size_pt is not None:
                run.font.size = Pt(float(font_size_pt))
            if bold is not None:
                run.bold = bool(bold)
            if italic is not None:
                run.italic = bool(italic)
            if underline is not None:
                run.underline = bool(underline)

    @staticmethod
    def _format_paragraph(dp: Any, cfg: dict[str, Any], *, is_listing: bool) -> None:
        # cfg: font_name/font_size_pt/line_spacing/first_line_indent_cm/alignment
        font_name = cfg.get("font_name")
        font_size_pt = cfg.get("font_size_pt")
        line_spacing = cfg.get("line_spacing")
        indent_cm = cfg.get("first_line_indent_cm")
        left_indent_cm = cfg.get("left_indent_cm")
        right_indent_cm = cfg.get("right_indent_cm")
        space_before_pt = cfg.get("space_before_pt")
        space_after_pt = cfg.get("space_after_pt")
        alignment = cfg.get("alignment")

        for run in dp.runs:
            if font_name:
                run.font.name = font_name
                # Нужно для некоторых шрифтов, чтобы Word корректно проставлял “eastAsia”.
                try:
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
                except Exception:
                    pass
            if font_size_pt is not None:
                run.font.size = Pt(float(font_size_pt))
            # Для упрощения делаем “чистый” стиль: без underline/italics/bold.
            run.underline = False
            run.bold = False
            run.italic = False

        if alignment:
            a = str(alignment).upper()
            if "CENTER" in a:
                dp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif "LEFT" in a:
                dp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif "RIGHT" in a:
                dp.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif "JUSTIFY" in a:
                dp.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        if line_spacing is not None:
            dp.paragraph_format.line_spacing = float(line_spacing)

        if indent_cm is not None:
            dp.paragraph_format.first_line_indent = Cm(float(indent_cm) / 1.0)
        if left_indent_cm is not None:
            dp.paragraph_format.left_indent = Cm(float(left_indent_cm))
        if right_indent_cm is not None:
            dp.paragraph_format.right_indent = Cm(float(right_indent_cm))
        if space_before_pt is not None:
            dp.paragraph_format.space_before = Pt(float(space_before_pt))
        if space_after_pt is not None:
            dp.paragraph_format.space_after = Pt(float(space_after_pt))

    @staticmethod
    def _format_list_numbering(doc: Any, cfg: dict[str, Any]) -> None:
        font_name = cfg.get("font_name") or "Times New Roman"
        font_size_pt = float(cfg.get("font_size_pt", 14))
        bold = bool(cfg.get("bold", False))
        italic = bool(cfg.get("italic", False))
        underline = bool(cfg.get("underline", False))

        for style_name in cfg.get("style_names") or ["List Number"]:
            try:
                style = doc.styles[style_name]
            except Exception:
                continue
            style.font.name = font_name
            style.font.size = Pt(font_size_pt)
            style.font.bold = bold
            style.font.italic = italic
            style.font.underline = underline
            try:
                r_pr = style._element.get_or_add_rPr()
                r_fonts = r_pr.get_or_add_rFonts()
                for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                    r_fonts.set(qn(f"w:{attr}"), font_name)
            except Exception:
                pass

        try:
            numbering = doc.part.numbering_part.element
        except Exception:
            return

        for lvl in numbering.findall(".//" + qn("w:lvl")):
            num_fmt = lvl.find(qn("w:numFmt"))
            if num_fmt is not None and num_fmt.get(qn("w:val")) not in {None, "decimal"}:
                continue
            r_pr = lvl.find(qn("w:rPr"))
            if r_pr is None:
                r_pr = OxmlElement("w:rPr")
                lvl.append(r_pr)

            r_fonts = r_pr.find(qn("w:rFonts"))
            if r_fonts is None:
                r_fonts = OxmlElement("w:rFonts")
                r_pr.append(r_fonts)
            for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                r_fonts.set(qn(f"w:{attr}"), font_name)

            size = r_pr.find(qn("w:sz"))
            if size is None:
                size = OxmlElement("w:sz")
                r_pr.append(size)
            size.set(qn("w:val"), str(int(font_size_pt * 2)))

            size_cs = r_pr.find(qn("w:szCs"))
            if size_cs is None:
                size_cs = OxmlElement("w:szCs")
                r_pr.append(size_cs)
            size_cs.set(qn("w:val"), str(int(font_size_pt * 2)))

            for tag, value in (("w:b", bold), ("w:i", italic), ("w:u", underline)):
                el = r_pr.find(qn(tag))
                if el is None:
                    el = OxmlElement(tag)
                    r_pr.append(el)
                if tag == "w:u":
                    el.set(qn("w:val"), "single" if value else "none")
                else:
                    el.set(qn("w:val"), "true" if value else "false")

